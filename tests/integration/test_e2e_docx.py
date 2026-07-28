"""End-to-end test of the HTTP API for a docx upload.

Drives the full upload → poll → download loop with a real
``TestClient`` and the mock privacy client.
"""

from __future__ import annotations

import shutil
from collections.abc import Generator
from pathlib import Path

import pytest
from docx import Document
from fastapi.testclient import TestClient
from neironir.api.dependencies import get_privacy, get_storage
from neironir.main import create_app
from neironir.privacy.client import MockPrivacyFilterClient
from neironir.storage.local import LocalStorage


def _build_docx(paragraphs: list[str], target: Path) -> Path:
    """Write a real ``.docx`` with the given plain-text paragraphs."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    document.save(str(target))
    return target


@pytest.fixture
def client_and_storage(tmp_path: Path) -> Generator[tuple[TestClient, LocalStorage], None, None]:
    storage_dir = tmp_path / "storage"
    storage_dir.mkdir()
    storage = LocalStorage(storage_dir)
    privacy = MockPrivacyFilterClient()

    app = create_app()
    app.dependency_overrides[get_storage] = lambda: storage
    app.dependency_overrides[get_privacy] = lambda: privacy

    with TestClient(app) as client:
        yield client, storage

    shutil.rmtree(storage_dir, ignore_errors=True)


def test_e2e_docx_upload_poll_download(
    client_and_storage: tuple[TestClient, LocalStorage], tmp_path: Path
) -> None:
    """A docx file is redacted and downloadable end-to-end."""
    client, _storage = client_and_storage
    source = _build_docx(
        [
            "Email: alice@example.com",
            "Phone: +7 495 123-45-67",
        ],
        tmp_path / "source.docx",
    )

    upload_response = client.post(
        "/api/v1/documents/",
        files={"file": ("contract.docx", source.read_bytes(), "application/octet-stream")},
    )
    assert upload_response.status_code == 202
    job = upload_response.json()
    job_id = job["id"]
    assert job["source_filename"] == "contract.docx"
    assert job["source_ext"] == "docx"

    # Poll for completion.
    final_status: str | None = None
    for _ in range(20):
        status_response = client.get(f"/api/v1/documents/{job_id}")
        assert status_response.status_code == 200
        final_status = status_response.json()["status"]
        if final_status in {"completed", "failed"}:
            break
    assert final_status == "completed"

    # Download and verify.
    download_response = client.get(f"/api/v1/documents/{job_id}/download")
    assert download_response.status_code == 200
    cleaned_bytes = download_response.content
    # The response is a real .docx, not a text dump.
    assert cleaned_bytes[:4] == b"PK\x03\x04"

    # Persist and read back with python-docx.
    result_path = tmp_path / "result.docx"
    result_path.write_bytes(cleaned_bytes)
    paragraphs = [p.text for p in Document(str(result_path)).paragraphs]

    assert "alice@example.com" not in cleaned_bytes.decode("latin-1", errors="ignore")
    assert "<PRIVATE_EMAIL1>" in paragraphs[0]
    assert "<PRIVATE_PHONE1>" in paragraphs[1]
    # The file's content-type and disposition convey its identity.
    assert 'filename="contract.cleaned.docx"' in download_response.headers["content-disposition"]


def test_e2e_docx_without_sensitive_data_round_trips(
    client_and_storage: tuple[TestClient, LocalStorage], tmp_path: Path
) -> None:
    """A clean docx is preserved as-is (no placeholder, no formatting loss)."""
    client, _storage = client_and_storage
    paragraphs = ["Hello", "world", "no PII here"]
    source = _build_docx(paragraphs, tmp_path / "source.docx")

    upload_response = client.post(
        "/api/v1/documents/",
        files={"file": ("hello.docx", source.read_bytes(), "application/octet-stream")},
    )
    assert upload_response.status_code == 202
    job_id = upload_response.json()["id"]

    for _ in range(20):
        status_response = client.get(f"/api/v1/documents/{job_id}")
        if status_response.json()["status"] in {"completed", "failed"}:
            break
    assert status_response.json()["status"] == "completed"

    download_response = client.get(f"/api/v1/documents/{job_id}/download")
    assert download_response.status_code == 200
    result_path = tmp_path / "result.docx"
    result_path.write_bytes(download_response.content)
    result_paragraphs = [p.text for p in Document(str(result_path)).paragraphs]
    assert result_paragraphs == paragraphs
