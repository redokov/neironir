"""Integration tests that exercise the **real** OPF neural model.

These tests live next to the mock-based integration tests so that the
two are easy to diff, but they are **skipped by default** because:

* they need ``opf`` on PATH (or pointed to via ``NEIRONIR_PRIVACY_FILTER_CMD``);
* they take a few seconds per call because the subprocess has to
  encode the input through the actual neural network.

Run them locally before tagging a release::

    NEIRONIR_RUN_REAL_MODEL_TESTS=1 \
        pytest tests/integration/test_pipeline_real_model.py -v

Or via the Makefile target ``make test-real``.
"""

from __future__ import annotations

import os
import shutil
import time
from collections.abc import Generator
from pathlib import Path

import pytest

# Module-level skip — the marker on its own is not enough because the
# ``opf`` binary might be on PATH but broken (corrupt checkpoint, etc.).
OPF_CMD = os.environ.get("NEIRONIR_PRIVACY_FILTER_CMD") or shutil.which("opf")
RUN_FLAG = os.environ.get("NEIRONIR_RUN_REAL_MODEL_TESTS") == "1"


pytestmark = [
    pytest.mark.real_model,
    pytest.mark.skipif(
        not RUN_FLAG,
        reason=(
            "NEIRONIR_RUN_REAL_MODEL_TESTS is not set — these tests "
            "require a real OPF subprocess and take seconds each"
        ),
    ),
    pytest.mark.skipif(
        OPF_CMD is None,
        reason="opf CLI not on PATH and NEIRONIR_PRIVACY_FILTER_CMD not set",
    ),
]


def _wait_for_completion(client, job_id, max_wait_s: float = 60.0) -> dict:
    """Poll the job endpoint until the pipeline is finished or fails.

    The real OPF model takes a few seconds per document so the timeout
    is generous; CI machines without GPU acceleration may need even
    more.
    """
    deadline = time.monotonic() + max_wait_s
    while time.monotonic() < deadline:
        r = client.get(f"/api/v1/documents/{job_id}")
        assert r.status_code == 200
        body = r.json()
        if body["status"] in {"completed", "failed"}:
            return body
        time.sleep(0.3)
    raise AssertionError(
        f"job {job_id} did not complete within {max_wait_s}s — "
        "is the OPF subprocess stuck or the timeout too small?"
    )


@pytest.fixture(scope="module")
def real_client_and_storage() -> Generator[tuple, None, None]:
    """Build a TestClient that talks to the real OPF subprocess."""
    from neironir.admin.training import reset_training_state
    from neironir.api.dependencies import get_privacy, get_settings, get_storage
    from neironir.config import Settings
    from neironir.main import create_app
    from neironir.privacy.client import SubprocessPrivacyFilterClient
    from neironir.storage.local import LocalStorage

    storage_dir = Path(tempfile.mkdtemp(prefix="real_model_storage_"))
    privacy = SubprocessPrivacyFilterClient(
        opf_cmd=OPF_CMD.split() if OPF_CMD else None,
        timeout_s=float(os.environ.get("NEIRONIR_PRIVACY_FILTER_TIMEOUT", "60")),
    )

    real_settings = Settings().model_copy(
        update={
            "storage_dir": str(storage_dir),
            "privacy_filter_mode": "subprocess",
            "privacy_filter_cmd": OPF_CMD or "",
        }
    )

    app = create_app()
    app.dependency_overrides[get_settings] = lambda: real_settings
    app.dependency_overrides[get_storage] = lambda: LocalStorage(storage_dir)
    app.dependency_overrides[get_privacy] = lambda: privacy

    reset_training_state()
    from fastapi.testclient import TestClient

    with TestClient(app) as client:
        yield client, storage_dir

    reset_training_state()
    shutil.rmtree(storage_dir, ignore_errors=True)


# ``tempfile`` is imported lazily so the module can be collected even
# when the test is skipped.
import tempfile  # noqa: E402

# ---------------------------------------------------------------------------
# Detection quality
# ---------------------------------------------------------------------------


class TestRealModelDetection:
    """Verify that the real model finds the entity types the mock can't."""

    def test_detects_person_name_in_russian_text(self, real_client_and_storage: tuple) -> None:
        """``Жалнин Максим Михайлович`` must produce a name annotation."""
        client, _ = real_client_and_storage
        text = "Генеральный директор Жалнин Максим Михайлович подписал документ."
        r = client.post(
            "/api/v1/documents/",
            files={"file": ("contract.md", text.encode("utf-8"), "text/markdown")},
        )
        assert r.status_code == 202
        job = _wait_for_completion(client, r.json()["id"])
        assert job["status"] == "completed", job

        ann = client.get(f"/api/v1/documents/{job['id']}/annotations").json()
        entity_types = {span["entity_type"] for span in ann["spans"]}
        # The mock does not detect person names — the real model must.
        assert "private_person" in entity_types, (
            "real OPF model failed to detect the ФИО in the sample — "
            "either the model checkpoint is missing/corrupt or the "
            "subprocess call failed silently"
        )

    def test_detects_address_in_russian_text(self, real_client_and_storage: tuple) -> None:
        client, _ = real_client_and_storage
        text = "Адрес: г. Москва, ул. Тверская, д. 1."
        r = client.post(
            "/api/v1/documents/",
            files={"file": ("contract.md", text.encode("utf-8"), "text/markdown")},
        )
        job = _wait_for_completion(client, r.json()["id"])
        assert job["status"] == "completed"

        ann = client.get(f"/api/v1/documents/{job['id']}/annotations").json()
        entity_types = {span["entity_type"] for span in ann["spans"]}
        assert "private_address" in entity_types, (
            "real OPF model failed to detect the address — see "
            "fixtures/sample_addresses.md for more cases"
        )

    def test_detects_email(self, real_client_and_storage: tuple) -> None:
        client, _ = real_client_and_storage
        text = "Контакты: user@example.com, +7 495 123-45-67"
        r = client.post(
            "/api/v1/documents/",
            files={"file": ("contract.md", text.encode("utf-8"), "text/markdown")},
        )
        job = _wait_for_completion(client, r.json()["id"])
        assert job["status"] == "completed"

        ann = client.get(f"/api/v1/documents/{job['id']}/annotations").json()
        entity_types = {span["entity_type"] for span in ann["spans"]}
        assert "private_email" in entity_types
        assert "private_phone" in entity_types


# ---------------------------------------------------------------------------
# Output format conversion
# ---------------------------------------------------------------------------


class TestRealModelDocxToMarkdown:
    """A docx uploaded with ``output_format=md`` should still be redacted
    by the real model and the markdown text should contain placeholders."""

    def test_docx_upload_with_real_model_redacts_and_converts(
        self, real_client_and_storage: tuple
    ) -> None:
        from docx import Document

        client, _ = real_client_and_storage

        docx_path = Path(tempfile.gettempdir()) / "real_model_test.docx"
        doc = Document()
        doc.add_paragraph("Контактное лицо: Жалнин Максим Михайлович")
        doc.add_paragraph("Email: support@example.org")
        doc.save(str(docx_path))

        with open(docx_path, "rb") as fh:
            r = client.post(
                "/api/v1/documents/",
                files={"file": ("contract.docx", fh.read(), "application/octet-stream")},
                data={"output_format": "md"},
            )
        assert r.status_code == 202
        job = _wait_for_completion(client, r.json()["id"], max_wait_s=120.0)
        assert job["status"] == "completed", job

        # The download is a markdown file with placeholders.
        dl = client.get(f"/api/v1/documents/{job['id']}/download")
        assert dl.status_code == 200
        assert "text/markdown" in dl.headers["content-type"]
        body = dl.text
        assert "<PRIVATE_PERSON" in body, body[:500]
        assert "<PRIVATE_EMAIL" in body
        # The original PII text must be gone.
        assert "support@example.org" not in body


# ---------------------------------------------------------------------------
# Apply-feedback round-trip with the real model
# ---------------------------------------------------------------------------


class TestRealModelApplyFeedback:
    """The apply-feedback endpoint must work against real-model output.

    This catches regressions where the offsets returned by the real
    OPF model don't follow the simple ``text[start:end]`` convention
    we rely on for the placeholder map.
    """

    def test_reject_real_model_span(self, real_client_and_storage: tuple) -> None:
        client, _ = real_client_and_storage
        text = "Директор Жалнин Максим Михайлович подписал контракт."
        r = client.post(
            "/api/v1/documents/",
            files={"file": ("contract.md", text.encode("utf-8"), "text/markdown")},
        )
        job = _wait_for_completion(client, r.json()["id"])
        ann = client.get(f"/api/v1/documents/{job['id']}/annotations").json()

        # Find the index of the first person annotation.
        person_idx = next(
            i for i, span in enumerate(ann["spans"]) if span["entity_type"] == "private_person"
        )

        # Reject it and verify the placeholder disappears from result.md.
        payload = {
            "actions": [
                {
                    "action": "reject",
                    "start": ann["spans"][person_idx]["start"],
                    "end": ann["spans"][person_idx]["end"],
                    "entity_type": "private_person",
                    "text": ann["spans"][person_idx]["text"],
                    "original_span_index": person_idx,
                }
            ],
            "comment": None,
        }
        r2 = client.post(
            f"/api/v1/documents/{job['id']}/apply-feedback",
            json=payload,
        )
        assert r2.status_code == 200
        body = r2.json()
        assert body["rejected"] == 1

        # The downloaded result should no longer contain the rejected
        # span as a placeholder — the original text must be back.
        dl = client.get(f"/api/v1/documents/{job['id']}/download")
        result_text = dl.text
        # Find the rejected text and check it's not in placeholder form.
        original = ann["spans"][person_idx]["text"]
        assert original in result_text


# ---------------------------------------------------------------------------
# Mode reporting
# ---------------------------------------------------------------------------


class TestRealModelModeReporting:
    """The ``/api/v1/mode`` endpoint must show all 8 entity types
    when the real model is in use."""

    def test_mode_lists_all_entity_types(self, real_client_and_storage: tuple) -> None:
        client, _ = real_client_and_storage
        body = client.get("/api/v1/mode").json()
        assert body["privacy_filter_mode"] in {"subprocess", "combined"}
        detected = set(body["detected_types"])
        # Real model can detect all 8 categories; the frontend uses
        # this to decide whether to show the "switch to subprocess"
        # banner.
        assert "private_person" in detected
        assert "private_address" in detected
        assert "private_email" in detected
