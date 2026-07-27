"""End-to-end pipeline test for the docx format."""

from __future__ import annotations

from pathlib import Path
from uuid import uuid4

from docx import Document
from neironir.config import Settings
from neironir.domain.job import Job, JobStatus
from neironir.privacy.client import MockPrivacyFilterClient
from neironir.storage.local import LocalStorage
from neironir.workers.pipeline import run_job


async def test_run_job_redacts_docx(tmp_path: Path) -> None:
    storage = LocalStorage(tmp_path)
    settings = Settings()
    privacy = MockPrivacyFilterClient()

    job_id = uuid4()
    job = Job(
        id=job_id,
        status=JobStatus.PENDING,
        source_filename="contract.docx",
        source_ext="docx",
    )
    storage.save_job(job)

    source_path = tmp_path / "contract.docx"
    document = Document()
    document.add_paragraph("My email is user@example.com.")
    document.add_paragraph("Backup: admin@example.org.")
    document.save(str(source_path))

    storage.save_source(job_id, "contract.docx", source_path.read_bytes())

    await run_job(job_id, settings=settings, storage=storage, privacy=privacy)

    final = storage.load_job(job_id)
    assert final.status == JobStatus.COMPLETED
    assert final.error is None

    result_path = storage.job_dir(job_id) / "result.docx"
    paragraphs = [p.text for p in Document(str(result_path)).paragraphs]
    assert paragraphs == [
        "My email is <PRIVATE_EMAIL1>.",
        "Backup: <PRIVATE_EMAIL2>.",
    ]
