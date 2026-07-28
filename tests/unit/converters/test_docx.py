"""Tests for :class:`neironir.converters.docx.DocxConverter`."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from neironir.converters.base import Replacement
from neironir.converters.docx import DocxConverter
from neironir.domain.entity_type import EntityType


def _make_docx(tmp_path: Path, paragraphs: list[str], name: str = "source.docx") -> Path:
    """Create a real ``.docx`` on disk with the given plain-text paragraphs."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    path = tmp_path / name
    document.save(str(path))
    return path


def test_extract_text_joins_paragraphs_with_newlines(tmp_path: Path) -> None:
    source = _make_docx(tmp_path, ["first", "second", "third"])
    converter = DocxConverter()

    assert converter.extract_text(source) == "first\nsecond\nthird"


def test_round_trip_with_replacement(tmp_path: Path) -> None:
    source = _make_docx(
        tmp_path,
        ["Email: user@example.com", "Plain paragraph"],
    )
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    # ``"Email: "`` is 7 chars, the email is 16 chars.
    replacement = Replacement(
        start=7,
        end=23,
        entity_type=EntityType.PRIVATE_EMAIL,
        placeholder="<PRIVATE_EMAIL1>",
    )
    converter.build(source, target, [replacement])

    rewritten = Document(str(target))
    paragraphs = [p.text for p in rewritten.paragraphs]
    assert paragraphs == ["Email: <PRIVATE_EMAIL1>", "Plain paragraph"]


def test_build_clips_replacement_at_element_boundary(tmp_path: Path) -> None:
    """A replacement that straddles a paragraph boundary is clipped to the first paragraph."""
    source = _make_docx(tmp_path, ["Email me at user@example.", "com soon"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    # The email starts in the first paragraph at offset 12 and ends in
    # the second one at offset 3 — the span is clipped to paragraph 0.
    # Note: text[12:] = "user@example." (14 chars). The replacement
    # asks for end=29, but paragraph 0 only has 26 chars (+"\n").
    # Since we clip, the placeholder replaces "user@example." in p0.
    replacement = Replacement(
        start=12,
        end=29,
        entity_type=EntityType.PRIVATE_EMAIL,
        placeholder="<PRIVATE_EMAIL1>",
    )
    converter.build(source, target, [replacement])

    rewritten = Document(str(target))
    paragraphs = [p.text for p in rewritten.paragraphs]
    # Paragraph 0 gets the placeholder, paragraph 1 is untouched
    # because the tail of the clipped span is dropped.
    assert paragraphs[0] == "Email me at <PRIVATE_EMAIL1>"


def test_build_with_no_replacements_preserves_paragraphs(tmp_path: Path) -> None:
    source = _make_docx(tmp_path, ["one", "two"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    converter.build(source, target, [])

    paragraphs = [p.text for p in Document(str(target)).paragraphs]
    assert paragraphs == ["one", "two"]
