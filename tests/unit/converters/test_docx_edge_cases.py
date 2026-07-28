"""Edge-case tests for :class:`neironir.converters.docx.DocxConverter`.

The basic happy-path tests live in :mod:`tests.unit.converters.test_docx`.
This file covers the corners: empty documents, single paragraphs, and
multi-paragraph documents where the cumulative-offset arithmetic could
trip up the converter.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from neironir.converters.base import Replacement
from neironir.converters.docx import DocxConverter
from neironir.domain.entity_type import EntityType


def _make_docx(tmp_path: Path, paragraphs: list[str], name: str = "src.docx") -> Path:
    """Create a real ``.docx`` with the given plain-text paragraphs."""
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    path = tmp_path / name
    document.save(str(path))
    return path


def _read_paragraphs(target: Path) -> list[str]:
    """Read the paragraphs of a ``.docx`` as plain strings."""
    return [paragraph.text for paragraph in Document(str(target)).paragraphs]


# ---------------------------------------------------------------------------
# extract_text
# ---------------------------------------------------------------------------


def test_extract_text_for_empty_document(tmp_path: Path) -> None:
    """A document with no paragraphs must extract to an empty string."""
    source = _make_docx(tmp_path, [])
    converter = DocxConverter()

    assert converter.extract_text(source) == ""


def test_extract_text_for_single_paragraph(tmp_path: Path) -> None:
    """A single-paragraph document has no separators to add."""
    source = _make_docx(tmp_path, ["only one"])
    converter = DocxConverter()

    assert converter.extract_text(source) == "only one"


def test_extract_text_preserves_empty_paragraphs_as_empty_segments(tmp_path: Path) -> None:
    """Empty paragraphs become empty segments between separators."""
    source = _make_docx(tmp_path, ["first", "", "third"])
    converter = DocxConverter()

    # The python-docx ``add_paragraph`` with an empty string still
    # produces a paragraph; the extracted text is joined with ``\n``
    # so we get the empty middle segment.
    assert converter.extract_text(source) == "first\n\nthird"


# ---------------------------------------------------------------------------
# build
# ---------------------------------------------------------------------------


def test_build_empty_document_round_trips(tmp_path: Path) -> None:
    """A round-trip on an empty document must produce an empty document."""
    source = _make_docx(tmp_path, [])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    converter.build(source, target, [])

    assert _read_paragraphs(target) == []


def test_build_single_paragraph_replacement(tmp_path: Path) -> None:
    """A replacement inside a single paragraph must work end-to-end."""
    source = _make_docx(tmp_path, ["Contact alice@example.com today"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    # The email is 17 chars and starts at offset 8.
    replacement = Replacement(
        start=8,
        end=25,
        entity_type=EntityType.PRIVATE_EMAIL,
        placeholder="<PRIVATE_EMAIL1>",
    )
    converter.build(source, target, [replacement])

    assert _read_paragraphs(target) == ["Contact <PRIVATE_EMAIL1> today"]


def test_build_replacements_in_multiple_paragraphs(tmp_path: Path) -> None:
    """Each replacement is anchored in its own paragraph."""
    source = _make_docx(tmp_path, ["Email: a@example.com", "Phone: +7 999 123 45 67"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    # Concatenated text: "Email: a@example.com\nPhone: +7 999 123 45 67"
    #  - email starts at 7, ends at 20
    #  - phone starts at 28, ends at 44
    replacements = [
        Replacement(7, 20, EntityType.PRIVATE_EMAIL, "<PRIVATE_EMAIL1>"),
        Replacement(28, 44, EntityType.PRIVATE_PHONE, "<PRIVATE_PHONE1>"),
    ]
    converter.build(source, target, replacements)

    assert _read_paragraphs(target) == [
        "Email: <PRIVATE_EMAIL1>",
        "Phone: <PRIVATE_PHONE1>",
    ]


def test_build_replacement_at_first_paragraph_offset_zero(tmp_path: Path) -> None:
    """A replacement at offset 0 in the first paragraph must work."""
    source = _make_docx(tmp_path, ["alice@example.com is mine", "second"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    # Email is 17 chars, starts at 0.
    replacement = Replacement(
        start=0,
        end=17,
        entity_type=EntityType.PRIVATE_EMAIL,
        placeholder="<PRIVATE_EMAIL1>",
    )
    converter.build(source, target, [replacement])

    assert _read_paragraphs(target) == ["<PRIVATE_EMAIL1> is mine", "second"]


def test_build_replacement_at_very_end_of_last_paragraph(tmp_path: Path) -> None:
    """A replacement ending at the very last character of the last paragraph works."""
    source = _make_docx(tmp_path, ["a", "tail@example.com"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    # Concatenated: "a\ntail@example.com" (length 2 + 1 + 15 = 18).
    # Email starts at 2, ends at 18 (== total length).
    replacement = Replacement(
        start=2,
        end=18,
        entity_type=EntityType.PRIVATE_EMAIL,
        placeholder="<PRIVATE_EMAIL1>",
    )
    converter.build(source, target, [replacement])

    assert _read_paragraphs(target) == ["a", "<PRIVATE_EMAIL1>"]


def test_build_replacement_crosses_boundary_after_empty_paragraph(tmp_path: Path) -> None:
    """A replacement crossing a boundary through an empty paragraph is clipped.

    The placeholder replaces text only in the first element; the rest
    of the span is dropped (no crash).
    """
    source = _make_docx(tmp_path, ["first", "", "second"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    replacement = Replacement(
        start=4,
        end=9,
        entity_type=EntityType.PRIVATE_EMAIL,
        placeholder="<PRIVATE_EMAIL1>",
    )
    converter.build(source, target, [replacement])

    # The span [4:9] starts in "first" at char 4..5 ("t") and crosses
    # into the empty paragraph. After clipping, only "t" in "first"
    # is replaced. The empty paragraph and "second" are untouched.
    rewritten = Document(str(target))
    paragraphs = [p.text for p in rewritten.paragraphs]
    assert paragraphs[0] == "firs<PRIVATE_EMAIL1>"


def test_build_replacements_supplied_in_reverse_order(tmp_path: Path) -> None:
    """The converter must accept replacements in any order."""
    source = _make_docx(tmp_path, ["a@b.com and c@d.com and e@f.com"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    # Three emails at positions (0, 7), (12, 19), (24, 31) in the
    # concatenated text. The pipeline is expected to pass them in
    # ascending order, but the converter must also accept the reverse
    # order — the spec contract says implementations are tolerant to
    # any input order.
    replacements = [
        Replacement(24, 31, EntityType.PRIVATE_EMAIL, "<PRIVATE_EMAIL3>"),
        Replacement(12, 19, EntityType.PRIVATE_EMAIL, "<PRIVATE_EMAIL2>"),
        Replacement(0, 7, EntityType.PRIVATE_EMAIL, "<PRIVATE_EMAIL1>"),
    ]
    converter.build(source, target, replacements)

    assert _read_paragraphs(target) == [
        "<PRIVATE_EMAIL1> and <PRIVATE_EMAIL2> and <PRIVATE_EMAIL3>"
    ]


def test_build_zero_length_replacement_at_offset(tmp_path: Path) -> None:
    """A zero-length replacement inserts the placeholder at a position."""
    source = _make_docx(tmp_path, ["ab"])
    converter = DocxConverter()
    target = tmp_path / "out.docx"

    replacement = Replacement(
        start=1,
        end=1,
        entity_type=EntityType.PRIVATE_EMAIL,
        placeholder="<PRIVATE_EMAIL1>",
    )
    converter.build(source, target, [replacement])

    assert _read_paragraphs(target) == ["a<PRIVATE_EMAIL1>b"]
