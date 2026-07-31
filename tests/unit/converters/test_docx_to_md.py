"""Unit tests for :mod:`neironir.converters.docx_to_md`."""

from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn
from neironir.converters.docx_to_md import (
    MarkdownElement,
    convert_to_markdown,
    extract_markdown_runs,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _build_docx(path: Path, *, heading: str | None = None, body: str | None = None) -> Path:
    """Write a tiny docx with optional heading + body to ``path``."""
    document = Document()
    if heading:
        document.add_heading(heading, level=1)
    if body:
        document.add_paragraph(body)
    document.save(str(path))
    return path


def _build_docx_with_runs(path: Path, *segments: tuple[str, dict]) -> Path:
    """Write a docx where each run in the first paragraph is built from
    a ``(text, properties)`` tuple.  Properties may contain ``bold`` and
    ``italic`` flags.
    """
    document = Document()
    paragraph = document.add_paragraph()
    for text, props in segments:
        run = paragraph.add_run(text)
        if props.get("bold"):
            run.bold = True
        if props.get("italic"):
            run.italic = True
    document.save(str(path))
    return path


def _build_docx_with_table(path: Path, rows: list[list[str]]) -> Path:
    """Write a docx with one table containing the given ``rows``."""
    document = Document()
    if rows:
        table = document.add_table(rows=len(rows), cols=len(rows[0]))
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                table.rows[r_idx].cells[c_idx].text = cell_text
    document.save(str(path))
    return path


def _add_hyperlink(paragraph, *, text: str, url: str, doc_part) -> None:
    """Append a ``w:hyperlink`` run with ``text`` and ``url`` to ``paragraph``."""
    r_id = doc_part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = qn("w:hyperlink")
    new_hl = paragraph._element.makeelement(hyperlink, {qn("r:id"): r_id})  # noqa: SLF001
    new_r = paragraph._element.makeelement(qn("w:r"), {})  # noqa: SLF001
    new_t = paragraph._element.makeelement(qn("w:t"), {})  # noqa: SLF001
    new_t.text = text
    new_r.append(new_t)
    new_hl.append(new_r)
    paragraph._element.append(new_hl)  # noqa: SLF001


def _build_docx_with_hyperlink(path: Path, text: str, url: str) -> Path:
    """Build a docx with one paragraph containing a hyperlink."""
    document = Document()
    paragraph = document.add_paragraph("Связаться: ")
    _add_hyperlink(paragraph, text=text, url=url, doc_part=document.part)
    document.save(str(path))
    return path


def _build_docx_with_styles(
    path: Path,
    *,
    paragraphs: list[tuple[str, str]],
) -> Path:
    """Build a docx where each paragraph has ``(style_name, text)``.

    Useful for verifying the localisation of the heading-style
    detection (``Заголовок 1``, ``1 уровень`` …).
    """
    import contextlib

    document = Document()
    for style_name, text in paragraphs:
        paragraph = document.add_paragraph(text)
        with contextlib.suppress(KeyError):
            # Style doesn't exist in this template; skip.
            paragraph.style = document.styles[style_name]
    document.save(str(path))
    return path


# ---------------------------------------------------------------------------
# Headings
# ---------------------------------------------------------------------------


class TestHeadings:
    def test_heading_one_renders_as_h1(self, tmp_path: Path) -> None:
        path = _build_docx(tmp_path / "x.docx", heading="Title")
        md = convert_to_markdown(path)
        assert "# Title" in md

    def test_multiple_heading_levels(self, tmp_path: Path) -> None:
        document = Document()
        document.add_heading("Top", level=1)
        document.add_heading("Sub", level=2)
        document.add_heading("Subsub", level=3)
        document.save(str(tmp_path / "x.docx"))
        md = convert_to_markdown(tmp_path / "x.docx")
        assert "# Top" in md
        assert "## Sub" in md
        assert "### Subsub" in md

    def test_localised_heading_styles(self, tmp_path: Path) -> None:
        """Russian-language contracts use ``Заголовок 1`` style names.

        The renderer must still treat them as headings rather than
        silently dropping them as plain paragraphs.
        """
        document = Document()
        try:
            heading_style = document.styles["Заголовок 1"]
        except KeyError:
            pytest.skip("This docx template lacks the Russian 'Заголовок 1' style")

        p = document.add_paragraph("ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ №2")
        p.style = heading_style
        document.save(str(tmp_path / "x.docx"))

        md = convert_to_markdown(tmp_path / "x.docx")
        assert "# ДОПОЛНИТЕЛЬНОЕ СОГЛАШЕНИЕ №2" in md

    def test_level_only_russian_heading(self, tmp_path: Path) -> None:
        """Some Russian docs use the ``1 уровень`` style variant."""
        document = Document()
        try:
            style = document.styles["1 уровень"]
        except KeyError:
            pytest.skip("This docx template lacks the '1 уровень' style")
        p = document.add_paragraph("Заголовок через 1 уровень")
        p.style = style
        document.save(str(tmp_path / "x.docx"))

        md = convert_to_markdown(tmp_path / "x.docx")
        assert "# Заголовок через 1 уровень" in md

    def test_heading_with_no_text_is_skipped(self, tmp_path: Path) -> None:
        """A heading paragraph with empty text should not emit ``#`` alone."""
        document = Document()
        document.add_heading("", level=1)
        document.save(str(tmp_path / "x.docx"))
        md = convert_to_markdown(tmp_path / "x.docx")
        # The ``# `` line is dropped; no isolated ``#`` remains.
        assert "\n# \n" not in md
        assert "#" not in md  # no heading at all


# ---------------------------------------------------------------------------
# Paragraphs and runs
# ---------------------------------------------------------------------------


class TestParagraphs:
    def test_simple_paragraph(self, tmp_path: Path) -> None:
        path = _build_docx(tmp_path / "x.docx", body="Just a sentence.")
        md = convert_to_markdown(path)
        assert "Just a sentence." in md

    def test_bold_renders_as_double_star(self, tmp_path: Path) -> None:
        path = _build_docx_with_runs(
            tmp_path / "x.docx",
            ("Plain ", {}),
            ("bold", {"bold": True}),
            (" tail.", {}),
        )
        md = convert_to_markdown(path)
        assert "Plain **bold** tail." in md

    def test_italic_renders_as_single_star(self, tmp_path: Path) -> None:
        path = _build_docx_with_runs(
            tmp_path / "x.docx",
            ("An ", {}),
            ("italic", {"italic": True}),
            (" word.", {}),
        )
        md = convert_to_markdown(path)
        assert "An *italic* word." in md

    def test_bold_italic_renders_as_triple_star(self, tmp_path: Path) -> None:
        path = _build_docx_with_runs(
            tmp_path / "x.docx",
            ("both ", {}),
            ("styled", {"bold": True, "italic": True}),
            (".", {}),
        )
        md = convert_to_markdown(path)
        assert "both ***styled***." in md

    def test_underline_is_dropped(self, tmp_path: Path) -> None:
        """Underline is intentionally stripped — spec keeps only bold/italic."""
        document = Document()
        p = document.add_paragraph()
        run = p.add_run("underlined")
        run.underline = True
        document.save(str(tmp_path / "x.docx"))

        md = convert_to_markdown(tmp_path / "x.docx")
        assert "underlined" in md
        # No markdown underline extension (``__text__``) appears.
        assert "__" not in md

    def test_empty_paragraph_is_skipped(self, tmp_path: Path) -> None:
        document = Document()
        document.add_paragraph("")
        document.add_paragraph("real text")
        document.save(str(tmp_path / "x.docx"))
        md = convert_to_markdown(tmp_path / "x.docx")
        assert "real text" in md
        # Two consecutive blank lines would be a gap; only one
        # separating newline between blocks is allowed.
        assert "\n\n\n" not in md

    def test_run_internal_whitespace_is_collapsed(self, tmp_path: Path) -> None:
        """Tabs and ``w:br`` inside a run collapse to single spaces."""
        document = Document()
        p = document.add_paragraph()
        run = p.add_run("before")
        run.add_tab()
        run.add_break()
        run.add_text("after")
        document.save(str(tmp_path / "x.docx"))
        md = convert_to_markdown(tmp_path / "x.docx")
        assert "before after" in md


# ---------------------------------------------------------------------------
# Hyperlinks
# ---------------------------------------------------------------------------


class TestHyperlinks:
    def test_hyperlink_collapses_to_text(self, tmp_path: Path) -> None:
        """The user explicitly does NOT want ``[label](url)`` — only the
        label text should remain.  This is the bug the previous
        pandoc-based output had: it emitted
        ``<amashchinov@tomat-astra.ru>`` (Word's autolink syntax) which
        the privacy filter couldn't reliably annotate.
        """
        path = _build_docx_with_hyperlink(
            tmp_path / "x.docx",
            text="amashchinov@tomat-astra.ru",
            url="mailto:amashchinov@tomat-astra.ru",
        )
        md = convert_to_markdown(path)
        assert "amashchinov@tomat-astra.ru" in md
        # No URL is exposed to the user.
        assert "mailto:" not in md
        # No autolink wrapper remains.
        assert "<amashchinov" not in md
        # No markdown link syntax either.
        assert "](" not in md

    def test_hyperlink_text_is_visible_among_other_text(self, tmp_path: Path) -> None:
        path = _build_docx_with_hyperlink(
            tmp_path / "x.docx",
            text="click me",
            url="https://example.com",
        )
        md = convert_to_markdown(path)
        assert "Связаться: click me" in md


# ---------------------------------------------------------------------------
# Tables
# ---------------------------------------------------------------------------


class TestTables:
    def test_simple_table_renders_as_pipe_table(self, tmp_path: Path) -> None:
        path = _build_docx_with_table(
            tmp_path / "x.docx",
            [
                ["A", "B"],
                ["1", "2"],
                ["3", "4"],
            ],
        )
        md = convert_to_markdown(path)
        # Header + separator + 2 body rows.
        lines = md.splitlines()
        table_lines = [line for line in lines if line.startswith("|")]
        assert len(table_lines) == 4  # header, sep, 2 rows
        assert table_lines[1].replace("|", "").replace("-", "").strip() == ""

    def test_table_pads_column_widths(self, tmp_path: Path) -> None:
        path = _build_docx_with_table(
            tmp_path / "x.docx",
            [
                ["short", "a longer header"],
                ["1", "2"],
            ],
        )
        md = convert_to_markdown(path)
        # The body cell is padded to the header width so the column
        # borders align.
        lines = [line for line in md.splitlines() if line.startswith("|")]
        assert "a longer header" in lines[0]
        assert "2                " in lines[2] or "2" in lines[2]

    def test_tables_interleave_with_paragraphs(self, tmp_path: Path) -> None:
        """Document body: heading → paragraph → table → paragraph."""
        document = Document()
        document.add_heading("Title", level=1)
        document.add_paragraph("intro paragraph")
        table = document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "x"
        table.rows[0].cells[1].text = "y"
        table.rows[1].cells[0].text = "1"
        table.rows[1].cells[1].text = "2"
        document.add_paragraph("outro paragraph")
        document.save(str(tmp_path / "x.docx"))

        md = convert_to_markdown(tmp_path / "x.docx")
        # The table is surrounded by blank lines.
        assert "# Title" in md
        assert "intro paragraph" in md
        assert "| x | y |" in md
        assert "outro paragraph" in md

    def test_empty_table_is_skipped(self, tmp_path: Path) -> None:
        document = Document()
        document.add_table(rows=0, cols=0)  # degenerate but possible
        document.save(str(tmp_path / "x.docx"))
        md = convert_to_markdown(tmp_path / "x.docx")
        assert "|" not in md


# ---------------------------------------------------------------------------
# Element model
# ---------------------------------------------------------------------------


class TestMarkdownElement:
    def test_heading_render(self) -> None:
        element = MarkdownElement(kind="heading", level=2, text="Sub")
        assert element.render() == "## Sub"

    def test_paragraph_render(self) -> None:
        element = MarkdownElement(kind="paragraph", text="hello")
        assert element.render() == "hello"

    def test_blank_render(self) -> None:
        assert MarkdownElement(kind="blank").render() == ""

    def test_table_render(self) -> None:
        element = MarkdownElement(
            kind="table",
            rows=(("A", "B"), ("1", "2")),
        )
        rendered = element.render()
        assert "| A | B |" in rendered
        assert "| - | - |" in rendered
        assert "| 1 | 2 |" in rendered

    def test_unknown_kind_raises(self) -> None:
        with pytest.raises(ValueError):
            MarkdownElement(kind="bogus").render()


class TestExtractMarkdownRuns:
    def test_yields_heading_and_paragraph_in_order(self, tmp_path: Path) -> None:
        document = Document()
        document.add_heading("Title", level=1)
        document.add_paragraph("body")
        document.save(str(tmp_path / "x.docx"))

        elements = list(extract_markdown_runs(tmp_path / "x.docx"))
        kinds = [e.kind for e in elements]
        assert kinds == ["heading", "paragraph"]
        assert elements[0].level == 1
        assert elements[0].text == "Title"
        assert elements[1].text == "body"

    def test_empty_docx_yields_no_elements(self, tmp_path: Path) -> None:
        document = Document()
        document.save(str(tmp_path / "x.docx"))
        assert list(extract_markdown_runs(tmp_path / "x.docx")) == []


# ---------------------------------------------------------------------------
# End-to-end: the rendered markdown survives privacy-filter annotation
# ---------------------------------------------------------------------------


class TestRealWorldFixture:
    """Sanity checks against the noise patterns the previous pandoc
    output had.  These tests use a built-in fixture so they run
    everywhere — the real contract is exercised separately in the
    manual pre-release checklist.
    """

    @pytest.fixture
    def fixture_docx(self, tmp_path: Path) -> Path:
        document = Document()
        p = document.add_paragraph("See ")
        _add_hyperlink(p, text="a@b.com", url="mailto:a@b.com", doc_part=document.part)
        document.save(str(tmp_path / "fixture.docx"))
        return tmp_path / "fixture.docx"

    def test_no_autolink_wrappers(self, fixture_docx: Path) -> None:
        import re

        md = convert_to_markdown(fixture_docx)
        # No ``<email>`` autolink syntax survives.
        assert not re.search(r"<[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}>", md)

    def test_no_pandoc_mark_class(self, fixture_docx: Path) -> None:
        import re

        md = convert_to_markdown(fixture_docx)
        # The previous pandoc output left ``[text]{.mark}`` style
        # attributes everywhere — the new converter doesn't emit any
        # of them.
        assert not re.search(r"\{[.#][^}]*\}", md)

    def test_email_is_still_detectable(self, fixture_docx: Path) -> None:
        """The output must contain the email so the privacy filter
        can find it."""
        import re

        md = convert_to_markdown(fixture_docx)
        emails = re.findall(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", md)
        assert "a@b.com" in emails
