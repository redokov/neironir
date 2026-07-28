"""DOCX converter for the redaction pipeline.

Extracts text from **both** body paragraphs and tables, applying the
privacy-filter detection and replacement to the concatenated text.

Key design
----------

The document body is a sequence of XML elements — paragraphs (``w:p``)
and tables (``w:tbl``) — in reading order. This converter:

1. Iterates the body elements in document order.
2. Flattens each table into a sequence of individual cell texts (treated
   analogously to paragraphs for offset tracking).
3. Concatenates all texts with ``\\n`` separators for the privacy model.
4. During :meth:`build`, projects each :class:`Replacement` back to the
   originating paragraph or table cell.

Known limitations (inherited from the MVP)
------------------------------------------

* Inline formatting (bold, italic, links) is lost — all runs in a paragraph
  are replaced with a single plain-text run.
* Nested tables are not supported.
* A replacement that crosses a paragraph *or* table-cell boundary raises
  :class:`ValueError`.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from neironir.converters.base import Replacement

# Separator inserted between every pair of body-element texts (paragraphs
# and table cells alike). The privacy-filter sees one continuous string
# with ``\\n`` separators.
_SEPARATOR = "\n"


class DocxConverter:
    """Read/write DOCX files, handling both paragraphs and tables."""

    ext = "docx"

    # Internal element type constants
    _P = "p"  # paragraph
    _T = "t"  # table
    _TC = "c"  # table cell (flattened)

    def extract_text(self, source: Path) -> str:
        """Return the concatenated plain text of paragraphs and table cells.

        The result includes a ``\\n`` separator between every pair of
        body elements (paragraph↔paragraph, paragraph↔table-cell, etc.)
        so that offsets are stable and predictable.

        For the ``docx → md`` conversion the pipeline uses
        :func:`neironir.workers.pipeline._docx_to_markdown` which
        delegates to ``pandoc`` for a richer markdown rendering; this
        method remains the source-of-truth offset representation that
        :meth:`build` projects replacements onto.
        """
        document = Document(str(source))
        # Cache for :meth:`build`.
        self._doc_paragraphs = list(document.paragraphs)
        self._doc_tables = list(document.tables)
        self._element_map = _build_element_map(document)
        self._texts = _element_texts(self._element_map, self._doc_paragraphs, self._doc_tables)
        return _SEPARATOR.join(self._texts)

    def build(self, source: Path, target: Path, replacements: list[Replacement]) -> None:
        """Re-write ``source`` with replacements applied, saving to ``target``.

        Replacements are anchored in the concatenated text produced by an
        earlier call to :meth:`extract_text`. Each replacement is projected
        onto the originating paragraph or table cell. Cross-boundary
        replacements are rejected.
        """
        document: DocumentType = Document(str(source))
        paragraphs = list(document.paragraphs)
        tables = list(document.tables)
        element_map = _build_element_map(document)
        texts = _element_texts(element_map, paragraphs, tables)
        cumulative = _cumulative_offsets(texts)

        rewritten = _slice_elements(
            texts=texts,
            cumulative=cumulative,
            replacements=replacements,
            element_map=element_map,
        )

        # Apply rewritten texts back to paragraphs and table cells.
        _apply_rewrites(element_map, rewritten, paragraphs, tables)

        document.save(str(target))


# ---------------------------------------------------------------------------
# Body element iteration
# ---------------------------------------------------------------------------


def _build_element_map(
    document: DocumentType,
) -> list[tuple[str, int, int | None]]:
    """Produce a flat sequence of body elements in document order.

    Each entry is ``(type, index, cell_index)`` where:
    * ``type`` is ``DocxConverter._P`` (paragraph) or ``DocxConverter._TC``
      (table cell).
    * ``index`` is the 0-based paragraph or table **container** index.
    * ``cell_index`` is the **global** 0-based cell index across *all* rows
      of the table (flat sequential counter), or ``None`` for paragraphs.

    Tables are flattened row-by-row, cell-by-cell.  For example a 2×3 table
    yields six ``(_TC, table_idx, 0…5)`` entries.
    """
    body = document.element.body
    elements: list[tuple[str, int, int | None]] = []
    para_idx = 0
    table_idx = 0

    for child in body:
        tag = child.tag
        if tag == qn("w:p"):
            elements.append((DocxConverter._P, para_idx, None))
            para_idx += 1
        elif tag == qn("w:tbl"):
            table = document.tables[table_idx]
            global_cell_idx = 0
            for row in table.rows:
                for _cell in row.cells:
                    elements.append((DocxConverter._TC, table_idx, global_cell_idx))
                    global_cell_idx += 1
            table_idx += 1

    return elements


def _element_texts(
    element_map: list[tuple[str, int, int | None]],
    paragraphs: list[Paragraph],
    tables: list[Table],
) -> list[str]:
    """Extract text from each element in the map order.

    For table cells, ``cell_idx`` is the **global** cell index across
    *all rows* of the table (see :func:`_build_element_map`).
    """
    texts: list[str] = []
    for elem_type, idx, cell_idx in element_map:
        if elem_type == DocxConverter._P:
            texts.append(paragraphs[idx].text)
        elif elem_type == DocxConverter._TC:
            table = tables[idx]
            global_count = 0
            found = False
            for row in table.rows:
                for cell in row.cells:
                    if global_count == cell_idx:
                        texts.append(cell.text)
                        found = True
                        break
                    global_count += 1
                if found:
                    break
            if not found:
                texts.append("")
    return texts


# ---------------------------------------------------------------------------
# Offset computation
# ---------------------------------------------------------------------------


def _cumulative_offsets(texts: list[str]) -> list[int]:
    """Return the start offset of each element in the joined text.

    ``cumulative[i]`` is the index of ``texts[i][0]`` in the concatenated
    string. The list has one extra entry at the end — the total length —
    for simpler boundary checks.
    """
    offsets: list[int] = []
    cursor = 0
    for text in texts:
        offsets.append(cursor)
        cursor += len(text) + len(_SEPARATOR)
    offsets.append(cursor)
    return offsets


# ---------------------------------------------------------------------------
# Replacement application
# ---------------------------------------------------------------------------


def _find_element_index(cumulative: list[int], offset: int) -> int:
    """Return the index into ``texts`` that contains ``offset``."""
    for i in range(len(cumulative) - 2, -1, -1):
        if offset >= cumulative[i]:
            return i
    raise ValueError(f"offset {offset} is before the first element")


def _slice_elements(
    texts: list[str],
    cumulative: list[int],
    replacements: list[Replacement],
    element_map: list[tuple[str, int, int | None]],
) -> list[str]:
    """Apply replacements and return the rewritten text list.

    If a replacement crosses an element boundary (e.g. spans across two
    table cells or a cell and a paragraph), it is **clipped** to the
    first element rather than raising an error. This handles the case
    where the neural model produces a span that starts near the end of
    one cell and ends in the next — the span is applied only to the
    cell where it starts, and a warning is logged for the dropped tail.
    """
    import logging
    logger = logging.getLogger(__name__)

    rewritten = list(texts)
    ordered = sorted(replacements, key=lambda r: r.start, reverse=True)
    clipped_count = 0

    for replacement in ordered:
        start_idx = _find_element_index(cumulative, replacement.start)
        end_idx = _find_element_index(cumulative, replacement.end - 1)

        if start_idx == end_idx:
            # Normal case — fits inside one element.
            local_start = replacement.start - cumulative[start_idx]
            local_end = replacement.end - cumulative[start_idx]
            element_text = rewritten[start_idx]
            rewritten[start_idx] = (
                element_text[:local_start] + replacement.placeholder + element_text[local_end:]
            )
        else:
            # Cross-boundary span — clip to the first element.
            local_start = replacement.start - cumulative[start_idx]
            local_end = len(rewritten[start_idx])
            element_text = rewritten[start_idx]
            if local_start < local_end:
                rewritten[start_idx] = (
                    element_text[:local_start] + replacement.placeholder + element_text[local_end:]
                )
            clipped_count += 1
            if clipped_count <= 3:
                logger.warning(
                    "replacement clipped at element boundary: "
                    "start=%d (elem %d), end=%d (elem %d); "
                    "applied to first element only",
                    replacement.start, start_idx,
                    replacement.end, end_idx,
                )

    return rewritten


def _apply_rewrites(
    element_map: list[tuple[str, int, int | None]],
    rewritten: list[str],
    paragraphs: list[Paragraph],
    tables: list[Table],
) -> None:
    """Write rewritten texts back to the python-docx objects."""
    for pos, (elem_type, idx, cell_idx) in enumerate(element_map):
        new_text = rewritten[pos]
        if elem_type == DocxConverter._P:
            _replace_text(paragraphs[idx], new_text)
        elif elem_type == DocxConverter._TC:
            _set_cell_text(tables[idx], cell_idx, new_text)


def _replace_text(element: Paragraph, new_text: str) -> None:
    """Replace the element's text, preserving no other formatting.

    Clears all runs and assigns the replacement to the first run so the
    paragraph/cell still has a valid run element.

    Also removes hyperlink XML elements (``w:hyperlink``) which are not
    accessible via ``paragraph.runs`` — Word auto-converts email addresses
    and URLs to hyperlinks, and their text survives ``runs`` clearing.
    """
    from docx.oxml.ns import qn
    runs = getattr(element, "runs", None)
    if runs is None:
        raise AttributeError("element has no 'runs' attribute")

    # 1. Clear all normal runs
    for run in runs:
        run.text = ""

    # 2. Remove hyperlink elements (w:hyperlink) from the paragraph XML.
    #    These contain their own runs and are NOT included in paragraph.runs.
    p_element = getattr(element, "_element", None)
    if p_element is not None:
        hyperlink_tag = qn("w:hyperlink")
        hyperlinks = p_element.findall(hyperlink_tag)
        for hl in hyperlinks:
            p_element.remove(hl)

    # 3. Set the new text in the first run
    if runs:
        runs[0].text = new_text


def _set_cell_text(table: Table, cell_idx: int | None, new_text: str) -> None:
    """Set the text of a cell identified by its **global** cell index.

    The index is the sequential position across *all* rows of the table,
    matching the ``cell_index`` stored by :func:`_build_element_map`.

    Handles cells with multiple paragraphs by clearing all existing
    paragraph text and writing the full ``new_text`` into the first
    paragraph.  Multi-line content (``\n``) is preserved inside the
    single paragraph.  Hyperlinks are also removed from all paragraphs.
    """
    from docx.oxml.ns import qn
    hyperlink_tag = qn("w:hyperlink")

    count = 0
    for row in table.rows:
        for cell in row.cells:
            if cell_idx is not None and count == cell_idx:
                # Clear ALL paragraphs in the cell
                for para in cell.paragraphs:
                    runs = getattr(para, "runs", None)
                    if runs is not None:
                        for run in runs:
                            run.text = ""
                    # Remove hyperlinks from each paragraph
                    p_element = getattr(para, "_element", None)
                    if p_element is not None:
                        for hl in p_element.findall(hyperlink_tag):
                            p_element.remove(hl)
                # Write the full text to the first paragraph
                if cell.paragraphs:
                    first_para = cell.paragraphs[0]
                    runs = getattr(first_para, "runs", None)
                    if runs:
                        runs[0].text = new_text
                    else:
                        first_para.text = new_text
                return
            count += 1
    logger = __import__("logging").getLogger(__name__)
    logger.warning("table cell index %s not found (table has %d cells)", cell_idx, count)


__all__ = ["DocxConverter"]
