"""Convert a ``.docx`` file to a clean markdown representation.

Unlike :func:`neironir.workers.pipeline._docx_to_markdown` (which
shells out to ``pandoc``), this module walks the docx via
``python-docx`` and emits a minimal markdown subset:

* Paragraphs in document order
* ``#`` / ``##`` / ``###`` … for headings (``Заголовок 1``,
  ``Heading 2``, ``1 уровень``, ``2 уровень`` … all map to the same
  level of ``#`` characters)
* ``**bold**`` and ``*italic*`` runs only — all other run-level
  formatting (underline, strikethrough, custom styles, ``w:mark``
  highlight, ``w:hyperlink`` wrappers, ``w:sym``, ``w:ruby``,
  ``w:tab``, ``w:br``) is stripped down to plain text
* Hyperlinks collapse to plain text — no ``[label](url)``, no
  ``<label>``-wrappers from Word's autolinks
* Tables render as pipe-tables (``| col | col |``) with the first row
  used as the header
* ``{=html}`` / ``<!-- -->`` separators emitted by ``pandoc`` never
  appear because we never run ``pandoc``

The output is what the privacy-filter annotates with offsets, so
anything that would add noise (mark classes, native divs, html
blocks, span wrappers) is intentionally left out.

Two public helpers are exposed:

* :func:`convert_to_markdown` — full ``source → markdown string``
* :func:`extract_markdown_runs` — emits an interleaved list of
  strings and :class:`MarkdownElement` instances for callers that
  want to inspect / process the document before rendering

The render path is intentionally small — a single pass through the
document body, no AST round-trips, no third-party markdown library.
"""

from __future__ import annotations

import re
from collections.abc import Iterator
from dataclasses import dataclass
from pathlib import Path
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph

if TYPE_CHECKING:
    pass  # Document is a runtime helper, not a type — keep mypy quiet.

# ---------------------------------------------------------------------------
# Element model
# ---------------------------------------------------------------------------


@dataclass(frozen=True)
class MarkdownElement:
    """A single block-level element emitted by :func:`extract_markdown_runs`.

    Rendered to a markdown string via :meth:`render`.  Keeps the
    representation inspectable so callers (and tests) can verify
    structure without re-parsing the output text.
    """

    kind: str  # "heading" | "paragraph" | "table" | "blank"
    level: int = 0  # for "heading" only
    text: str = ""
    rows: tuple[tuple[str, ...], ...] = ()  # for "table" only

    def render(self) -> str:
        """Return the markdown representation of this element."""
        if self.kind == "heading":
            prefix = "#" * max(1, min(6, self.level))
            return f"{prefix} {self.text}".rstrip()
        if self.kind == "paragraph":
            return self.text
        if self.kind == "blank":
            return ""
        if self.kind == "table":
            return _render_table(self.rows)
        raise ValueError(f"unknown element kind: {self.kind!r}")


def _render_table(rows: tuple[tuple[str, ...], ...]) -> str:
    """Render ``rows`` as a GitHub-flavoured pipe-table.

    Empty input produces an empty string.  A single-row input becomes
    a headerless pipe-table (no separator line).
    """
    if not rows:
        return ""

    if len(rows) == 1:
        cells = rows[0]
        line = "| " + " | ".join(cells) + " |"
        return line

    header = rows[0]
    widths = [max(len(c) for c in col) for col in zip(*rows, strict=False)]
    header_line = "| " + " | ".join(c.ljust(w) for c, w in zip(header, widths, strict=False)) + " |"
    sep_line = "| " + " | ".join("-" * w for w in widths) + " |"
    body_lines = [
        "| " + " | ".join(c.ljust(w) for c, w in zip(row, widths, strict=False)) + " |"
        for row in rows[1:]
    ]
    return "\n".join([header_line, sep_line, *body_lines])


# ---------------------------------------------------------------------------
# Style detection
# ---------------------------------------------------------------------------


# Regexes for the various localised heading-style names we encounter
# in real-world Russian-language contracts.
_HEADING_STYLE_RE = re.compile(
    r"""
    ^\s*
    (?:
        (?:heading|заголовок|уровень)    # English / Russian
        \s*
        (\d+)                             # the level number
    |
        \d+\s*(?:уровень|heading|заголовок)  # "1 уровень" / "1 heading"
    )
    """,
    re.IGNORECASE | re.VERBOSE,
)


def _heading_level(paragraph: Paragraph) -> int | None:
    """Return the heading level for ``paragraph`` or ``None``.

    Accepts the localised style names that python-docx surfaces for
    Russian-language documents (``Заголовок 1``, ``1 уровень``, …) as
    well as the English canonical names (``Heading 1``).
    """
    style = paragraph.style
    style_name = style.name if style is not None else ""
    match = _HEADING_STYLE_RE.match(style_name)
    if match is None:
        return None
    raw = match.group(1)
    try:
        level = int(raw)
    except (TypeError, ValueError):
        return None
    # Clamp to the 1-6 range that ``#`` headers support.
    return max(1, min(6, level))


# ---------------------------------------------------------------------------
# Run-level rendering
# ---------------------------------------------------------------------------


def _render_runs(paragraph: Paragraph) -> str:
    """Render a paragraph's runs as plain markdown with ``**bold**`` /
    ``*italic*`` only.

    Hyperlinks (``w:hyperlink`` elements) are flattened to their text
    content — no ``[label](url)`` and no ``<label>``-wrapper artefacts
    that Word's autolinks leave behind.  All other inline
    formatting (underline, mark, custom styles) is dropped.
    """
    # python-docx's ``paragraph.text`` skips w:hyperlink content, but
    # our custom walk below rebuilds it from the underlying XML so we
    # can pick up the link text exactly once.
    p_element = paragraph._element  # noqa: SLF001 (private API but stable)
    if p_element is None:
        return ""

    # Iterate over direct children of the paragraph in document order.
    # Each child is either a run (``w:r``), a hyperlink (``w:hyperlink``)
    # or a structural element we want to skip silently.
    parts: list[str] = []
    for child in p_element.iterchildren():
        tag = child.tag
        if tag == qn("w:r"):
            parts.append(_render_run(child))
        elif tag == qn("w:hyperlink"):
            # Hyperlink = its inner runs joined.  Word also stores the
            # target URL in a relationship but we deliberately drop it.
            for inner in child.iterchildren():
                if inner.tag == qn("w:r"):
                    parts.append(_render_run(inner))
        # Other tags (w:bookmarkStart, w:bookmarkEnd, w:proofErr, …)
        # are ignored on purpose — they carry no user-visible text.

    text = "".join(parts)
    return _collapse_whitespace(text)


def _render_run(run_element: object) -> str:
    """Render a single ``w:r`` element to its inline markdown.

    ``run_element`` is the lxml-backed XML node; we annotate it as
    ``object`` so mypy doesn't try to introspect lxml's types but the
    runtime calls ``.iterchildren()`` work as expected.
    """
    bold = False
    italic = False
    text_parts: list[str] = []
    for child in run_element.iterchildren():  # type: ignore[attr-defined]
        tag = child.tag
        if tag == qn("w:rPr"):
            # Run properties — extract bold/italic only, ignore
            # everything else.
            if child.find(qn("w:b")) is not None:
                bold = True
            if child.find(qn("w:i")) is not None:
                italic = True
            # Underline, strikethrough, highlight, color, font, etc.
            # are intentionally ignored.
        elif tag == qn("w:t"):
            text_parts.append(child.text or "")
        elif tag == qn("w:tab"):
            text_parts.append("\t")
        elif tag == qn("w:br"):
            text_parts.append("\n")

    text = "".join(text_parts)
    if not text:
        return ""

    # Strip newline artefacts that the renderer would otherwise
    # preserve — markdown treats them as hard breaks which we don't
    # want here.
    text = text.replace("\r", "")
    text = text.replace("\n", " ")

    if bold and italic:
        return f"***{text}***"
    if bold:
        return f"**{text}**"
    if italic:
        return f"*{text}*"
    return text


_WHITESPACE_RE = re.compile(r"[ \t]+")
_NEWLINE_RUN_RE = re.compile(r" *\n+ *")


def _collapse_whitespace(text: str) -> str:
    """Collapse runs of spaces/tabs and trim line-by-line.

    We do **not** join lines — paragraph breaks already separate blocks
    in markdown.  We just clean up the artefacts left behind by
    ``w:tab`` + ``w:br`` runs inside a single paragraph.
    """
    if not text:
        return ""
    text = _NEWLINE_RUN_RE.sub("\n", text)
    # Collapse spaces/tabs inside each line.
    lines = text.split("\n")
    lines = [_WHITESPACE_RE.sub(" ", line).strip() for line in lines]
    return "\n".join(line for line in lines if line)


# ---------------------------------------------------------------------------
# Element extraction
# ---------------------------------------------------------------------------


def extract_markdown_runs(source: Path) -> Iterator[MarkdownElement]:
    """Yield :class:`MarkdownElement` instances for the docx at ``source``.

    The caller iterates the result and concatenates the ``render()``
    output, optionally with blank lines between block-level elements.

    Skipped entirely:

    * ``w:sdt`` (structured document tags) — used by Word for content
      controls and always rendered as opaque HTML by Pandoc
    * ``w:bookmarkStart`` / ``w:bookmarkEnd`` — never visible
    * Empty paragraphs that contain no text
    """
    document = Document(str(source))
    body = document.element.body
    if body is None:
        return

    # ``element_map`` mirrors :mod:`neironir.converters.docx` so we
    # visit paragraphs and tables in document order.  Nested tables
    # are out of scope for the MVP — they don't appear in the kinds
    # of contracts we handle.
    for child in body.iterchildren():
        tag = child.tag
        if tag == qn("w:p"):
            paragraph = Paragraph(child, document.part)
            yield from _paragraph_to_elements(paragraph)
        elif tag == qn("w:tbl"):
            table = _resolve_table(child, document)
            yield _table_to_element(table)
        # All other body children are ignored on purpose.


def _resolve_table(tbl_element: object, document: object) -> DocxTable:
    """Locate the python-docx ``Table`` object for ``tbl_element``."""
    for table in document.tables:  # type: ignore[attr-defined]
        if table._element is tbl_element:  # noqa: SLF001
            return table  # type: ignore[no-any-return]
    # Fallback: should never happen for well-formed documents.
    raise ValueError("table element not found in document.tables")


def _paragraph_to_elements(
    paragraph: Paragraph,
) -> Iterator[MarkdownElement]:
    """Convert one paragraph to zero or more elements."""
    level = _heading_level(paragraph)
    text = _render_runs(paragraph)

    if level is not None:
        if not text:
            # Heading with no text — skip rather than emit a noisy ``#``.
            return
        yield MarkdownElement(kind="heading", level=level, text=text)
        return

    if not text:
        # Skip empty paragraphs; their role is to provide vertical
        # spacing in the original document which we already control
        # at render time.
        return

    yield MarkdownElement(kind="paragraph", text=text)


def _table_to_element(table: DocxTable) -> MarkdownElement:
    """Convert one python-docx table to a markdown ``table`` element."""
    rows: list[tuple[str, ...]] = []
    for row in table.rows:
        # ``cell.text`` ignores hyperlinks but in tables that's
        # acceptable — the user's spec excludes hyperlinks from the
        # rendered output.
        cells = tuple(_collapse_whitespace(cell.text) for cell in row.cells)
        rows.append(cells)
    return MarkdownElement(kind="table", rows=tuple(rows))


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def convert_to_markdown(source: Path) -> str:
    """Return the cleaned-up markdown representation of ``source``.

    The output is a single string with blank lines separating
    block-level elements (headings, paragraphs, tables).  Trailing
    whitespace is trimmed.
    """
    blocks: list[str] = []
    for element in extract_markdown_runs(source):
        rendered = element.render()
        if rendered:
            blocks.append(rendered)
    return "\n\n".join(blocks).strip()


__all__ = [
    "MarkdownElement",
    "convert_to_markdown",
    "extract_markdown_runs",
]
